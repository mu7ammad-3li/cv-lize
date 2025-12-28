"""
DOCX generator for ATS-optimized resumes
Following strict ATS formatting guidelines from new-mods.md
"""

from io import BytesIO
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from models.schemas import ParsedResumeData


class DOCXGenerator:
    """
    Generate ATS-optimized DOCX resumes

    ATS Requirements:
    - Single-column linear layout
    - Standard fonts (Arial, Calibri)
    - Font sizes: Body 10-12pt, Headers 14-16pt, Name 24-36pt
    - 1-inch margins
    - Left-aligned text
    - 1.15-1.5 line spacing
    - Standard bullet points (• or -)
    - Contact info in main body (NOT in header/footer)
    - Standard section headers
    """

    # ATS-safe fonts
    SAFE_FONTS = {
        "sans_serif": ["Calibri", "Arial", "Helvetica", "Aptos"],
        "serif": ["Times New Roman", "Georgia", "Garamond", "Cambria"],
    }

    # Font sizes (in points)
    FONT_SIZES = {
        "name": 28,  # 24-36pt range
        "section_header": 14,  # 14-16pt range
        "body": 11,  # 10-12pt range
        "contact": 10,  # 10-12pt range
    }

    def __init__(self, font_family: str = "Calibri"):
        """
        Initialize DOCX generator

        Args:
            font_family: ATS-safe font (default: Calibri)
        """
        if font_family not in self.SAFE_FONTS["sans_serif"] + self.SAFE_FONTS["serif"]:
            print(
                f"⚠️  Warning: {font_family} may not be ATS-safe. Using Calibri instead."
            )
            font_family = "Calibri"

        self.font_family = font_family

    def generate_from_markdown(
        self, markdown_text: str, parsed_data: Optional[ParsedResumeData] = None
    ) -> bytes:
        """
        Generate DOCX from markdown resume

        Args:
            markdown_text: Optimized resume in markdown format
            parsed_data: Optional structured resume data

        Returns:
            DOCX file as bytes
        """
        doc = Document()
        self._setup_document_formatting(doc)

        # Parse markdown and build document
        self._parse_and_build(doc, markdown_text, parsed_data)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_from_parsed_resume(self, parsed_data: ParsedResumeData) -> bytes:
        """
        Generate DOCX from structured resume data

        Args:
            parsed_data: Structured resume data

        Returns:
            DOCX file as bytes
        """
        doc = Document()
        self._setup_document_formatting(doc)

        # Build document from structured data
        self._build_from_structured_data(doc, parsed_data)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _setup_document_formatting(self, doc: Document):
        """
        Set up ATS-compliant document formatting

        - 1-inch margins on all sides
        - Single-column layout
        - Standard spacing
        """
        # Set margins to 1 inch (914400 EMUs = 1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Single column (default)
            # Ensure no headers/footers are used for contact info
            section.different_first_page_header_footer = False

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        style: str = "body",
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """
        Add a paragraph with ATS-compliant formatting

        Args:
            doc: Document object
            text: Text content
            style: Paragraph style (name, section_header, body, contact)
            bold: Make text bold
            italic: Make text italic
        """
        p = doc.add_paragraph()
        run = p.add_run(text)

        # Set font
        run.font.name = self.font_family
        run.font.size = Pt(self.FONT_SIZES.get(style, self.FONT_SIZES["body"]))
        run.bold = bold
        run.italic = italic

        # Set paragraph formatting
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Always left-aligned

        # Line spacing: 1.08 (lighter, more compact)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.08

        # Reduced spacing after paragraphs for lighter appearance
        if style == "section_header":
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(6)
        elif style == "name":
            p.paragraph_format.space_after = Pt(2)
        else:
            p.paragraph_format.space_after = Pt(1)

        # Remove space before for tighter layout
        if style != "section_header":
            p.paragraph_format.space_before = Pt(0)

    def _add_bullet_point(self, doc: Document, text: str) -> None:
        """
        Add a bullet point with standard bullet character (•)

        Args:
            doc: Document object
            text: Bullet text
        """
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.name = self.font_family
        run.font.size = Pt(self.FONT_SIZES["body"])

        # Left-aligned
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Lighter line spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.08

        # Tighter spacing for bullets
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)

        # Indent
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    def _add_section_divider(self, doc: Document) -> None:
        """
        Add a subtle section divider (horizontal line)

        ATS-safe alternative to graphics/text boxes
        """
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 1/8 pt increments
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _parse_and_build(
        self, doc: Document, markdown_text: str, parsed_data: Optional[ParsedResumeData]
    ):
        """
        Parse markdown and build DOCX document

        Handles markdown headers, bullets, and formatting
        """
        lines = markdown_text.strip().split("\n")

        for i, line in enumerate(lines):
            line = line.strip()

            if not line:
                continue

            # H1 - Name (first H1 only)
            if line.startswith("# ") and i < 5:
                name = line[2:].strip()
                self._add_paragraph(doc, name, style="name", bold=True)

            # H2 - Section Headers
            elif line.startswith("## "):
                if i > 0:  # Add divider before sections (except first)
                    self._add_section_divider(doc)
                section_title = line[3:].strip()
                self._add_paragraph(
                    doc, section_title, style="section_header", bold=True
                )

            # H3 - Subsection (e.g., Job Title)
            elif line.startswith("### "):
                subsection = line[4:].strip()
                self._add_paragraph(doc, subsection, style="body", bold=True)

            # Bullet points
            elif line.startswith("- ") or line.startswith("* "):
                bullet_text = line[2:].strip()
                self._add_bullet_point(doc, bullet_text)

            # Italic dates (e.g., *MM/YYYY - MM/YYYY*)
            elif line.startswith("*") and line.endswith("*") and "/" in line:
                date_text = line.strip("*")
                self._add_paragraph(doc, date_text, style="body", italic=True)

            # Bold text (e.g., **Languages**)
            elif line.startswith("**") and line.endswith("**"):
                bold_text = line.strip("*")
                self._add_paragraph(doc, bold_text, style="body", bold=True)

            # Regular paragraph
            else:
                self._add_paragraph(doc, line, style="body")

    def _build_from_structured_data(self, doc: Document, parsed_data: ParsedResumeData):
        """
        Build DOCX from structured resume data

        Args:
            doc: Document object
            parsed_data: Structured resume data
        """
        # Name
        name = parsed_data.personalInfo.get("name", "Your Name")
        self._add_paragraph(doc, name, style="name", bold=True)

        # Contact Information (in main body, NOT header)
        contact_parts = []
        if parsed_data.personalInfo.get("email"):
            contact_parts.append(parsed_data.personalInfo["email"])
        if parsed_data.personalInfo.get("phone"):
            contact_parts.append(parsed_data.personalInfo["phone"])
        if parsed_data.personalInfo.get("location"):
            contact_parts.append(parsed_data.personalInfo["location"])
        if parsed_data.personalInfo.get("linkedin"):
            contact_parts.append(parsed_data.personalInfo["linkedin"])

        if contact_parts:
            self._add_paragraph(doc, " | ".join(contact_parts), style="contact")

        # Professional Summary
        if parsed_data.personalInfo.get("summary"):
            self._add_section_divider(doc)
            self._add_paragraph(
                doc, "Professional Summary", style="section_header", bold=True
            )
            self._add_paragraph(doc, parsed_data.personalInfo["summary"], style="body")

        # Technical Skills
        if parsed_data.skills:
            self._add_section_divider(doc)
            self._add_paragraph(
                doc, "Technical Skills", style="section_header", bold=True
            )

            for category, skills_list in parsed_data.skills.items():
                if skills_list:
                    category_name = category.replace("_", " ").title()
                    skills_str = (
                        ", ".join(skills_list)
                        if isinstance(skills_list, list)
                        else skills_list
                    )
                    # Create bullet point with bold category name
                    p = doc.add_paragraph(style="List Bullet")

                    # Bold category name
                    run_bold = p.add_run(f"{category_name}: ")
                    run_bold.font.name = self.font_family
                    run_bold.font.size = Pt(self.FONT_SIZES["body"])
                    run_bold.bold = True

                    # Normal skills list
                    run_normal = p.add_run(skills_str)
                    run_normal.font.name = self.font_family
                    run_normal.font.size = Pt(self.FONT_SIZES["body"])

                    # Apply formatting
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p.paragraph_format.line_spacing = 1.08
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)

        # Work Experience
        if parsed_data.experience:
            self._add_section_divider(doc)
            self._add_paragraph(
                doc, "Work Experience", style="section_header", bold=True
            )

            for exp in parsed_data.experience:
                # Job title and company
                title_company = (
                    f"{exp.get('title', 'Position')} | {exp.get('company', 'Company')}"
                )
                self._add_paragraph(doc, title_company, style="body", bold=True)

                # Dates
                if exp.get("duration") or (
                    exp.get("start_date") and exp.get("end_date")
                ):
                    dates = (
                        exp.get("duration")
                        or f"{exp.get('start_date')} - {exp.get('end_date')}"
                    )
                    self._add_paragraph(doc, dates, style="body", italic=True)

                # Responsibilities/achievements
                if exp.get("responsibilities"):
                    for resp in exp["responsibilities"]:
                        self._add_bullet_point(doc, resp)
                elif exp.get("description"):
                    # Handle description as either string or list
                    if isinstance(exp["description"], list):
                        for line in exp["description"]:
                            if line.strip():
                                self._add_bullet_point(doc, line.strip())
                    else:
                        # Split description string into bullet points if needed
                        desc_lines = exp["description"].split("\n")
                        for line in desc_lines:
                            if line.strip():
                                self._add_bullet_point(doc, line.strip())

                # Add minimal spacing between jobs
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)

        # Education
        if parsed_data.education:
            self._add_section_divider(doc)
            self._add_paragraph(doc, "Education", style="section_header", bold=True)

            for edu in (
                parsed_data.education.values()
                if isinstance(parsed_data.education, dict)
                else [parsed_data.education]
            ):
                if isinstance(edu, dict):
                    degree_institution = f"{edu.get('degree', 'Degree')} | {edu.get('institution', 'Institution')}"
                    self._add_paragraph(
                        doc, degree_institution, style="body", bold=True
                    )

                    if edu.get("year"):
                        self._add_paragraph(
                            doc, str(edu["year"]), style="body", italic=True
                        )

                    if edu.get("gpa"):
                        self._add_paragraph(doc, f"GPA: {edu['gpa']}", style="body")

        # Certifications
        if parsed_data.certifications:
            self._add_section_divider(doc)
            self._add_paragraph(
                doc, "Certifications", style="section_header", bold=True
            )
            for cert in parsed_data.certifications:
                self._add_bullet_point(doc, cert)


# Create singleton instance
docx_generator = DOCXGenerator()
